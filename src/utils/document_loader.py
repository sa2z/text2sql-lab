"""
문서 로더 유틸리티

다양한 형식의 문서(PDF, Word, Excel)를 로드하고 파싱하는 유틸리티
"""
import os
from typing import List, Dict, Any, Optional
from pathlib import Path
import logging

# PDF 처리
try:
    import PyPDF2
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# Word 문서 처리
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Excel 처리
try:
    import pandas as pd
    import openpyxl
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class Document:
    """문서 객체"""
    
    def __init__(self, page_content: str, metadata: Optional[Dict[str, Any]] = None):
        self.page_content = page_content
        self.metadata = metadata or {}
    
    def __repr__(self):
        return f"Document(content_length={len(self.page_content)}, metadata={self.metadata})"


class DocumentLoader:
    """다양한 형식의 문서를 로드하는 클래스"""
    
    def __init__(self):
        self.supported_formats = {
            '.pdf': self.load_pdf,
            '.docx': self.load_docx,
            '.doc': self.load_docx,
            '.xlsx': self.load_excel,
            '.xls': self.load_excel,
            '.txt': self.load_text
        }
    
    def load(self, file_path: str, **kwargs) -> List[Document]:
        """
        파일 확장자에 따라 적절한 로더를 자동으로 선택
        
        Args:
            file_path: 로드할 파일 경로
            **kwargs: 각 로더에 전달할 추가 인자
        
        Returns:
            Document 객체 리스트
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"파일을 찾을 수 없습니다: {file_path}")
        
        extension = path.suffix.lower()
        
        if extension not in self.supported_formats:
            raise ValueError(f"지원하지 않는 파일 형식입니다: {extension}")
        
        loader_func = self.supported_formats[extension]
        return loader_func(file_path, **kwargs)
    
    def load_pdf(self, file_path: str, use_pdfplumber: bool = True, **kwargs) -> List[Document]:
        """
        PDF 파일 로드
        
        Args:
            file_path: PDF 파일 경로
            use_pdfplumber: True면 pdfplumber 사용, False면 PyPDF2 사용
            **kwargs: 추가 옵션
        
        Returns:
            Document 객체 리스트 (페이지별)
        """
        if not PDF_AVAILABLE:
            raise ImportError("PDF 처리를 위해 PyPDF2와 pdfplumber를 설치해주세요: pip install PyPDF2 pdfplumber")
        
        documents = []
        file_name = Path(file_path).name
        
        try:
            if use_pdfplumber and pdfplumber:
                # pdfplumber 사용 (테이블 추출에 유리)
                with pdfplumber.open(file_path) as pdf:
                    for page_num, page in enumerate(pdf.pages, start=1):
                        text = page.extract_text()
                        if text and text.strip():
                            documents.append(Document(
                                page_content=text.strip(),
                                metadata={
                                    'source': file_path,
                                    'file_name': file_name,
                                    'page': page_num,
                                    'total_pages': len(pdf.pages),
                                    'type': 'pdf'
                                }
                            ))
            else:
                # PyPDF2 사용
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page_num, page in enumerate(pdf_reader.pages, start=1):
                        text = page.extract_text()
                        if text and text.strip():
                            documents.append(Document(
                                page_content=text.strip(),
                                metadata={
                                    'source': file_path,
                                    'file_name': file_name,
                                    'page': page_num,
                                    'total_pages': len(pdf_reader.pages),
                                    'type': 'pdf'
                                }
                            ))
            
            logger.info(f"PDF 로드 완료: {file_name} ({len(documents)} 페이지)")
            return documents
            
        except Exception as e:
            logger.error(f"PDF 로드 실패: {file_path}, 에러: {str(e)}")
            raise
    
    def load_docx(self, file_path: str, **kwargs) -> List[Document]:
        """
        Word 문서 로드
        
        Args:
            file_path: Word 파일 경로
            **kwargs: 추가 옵션
        
        Returns:
            Document 객체 리스트 (단락별)
        """
        if not DOCX_AVAILABLE:
            raise ImportError("Word 문서 처리를 위해 python-docx를 설치해주세요: pip install python-docx")
        
        documents = []
        file_name = Path(file_path).name
        
        try:
            doc = docx.Document(file_path)
            
            # 문단별로 처리
            for para_num, paragraph in enumerate(doc.paragraphs, start=1):
                text = paragraph.text.strip()
                if text:
                    documents.append(Document(
                        page_content=text,
                        metadata={
                            'source': file_path,
                            'file_name': file_name,
                            'paragraph': para_num,
                            'type': 'docx'
                        }
                    ))
            
            # 테이블도 추출
            for table_num, table in enumerate(doc.tables, start=1):
                table_text = []
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    table_text.append(' | '.join(row_text))
                
                if table_text:
                    documents.append(Document(
                        page_content='\n'.join(table_text),
                        metadata={
                            'source': file_path,
                            'file_name': file_name,
                            'table': table_num,
                            'type': 'docx_table'
                        }
                    ))
            
            logger.info(f"Word 문서 로드 완료: {file_name} ({len(documents)} 청크)")
            return documents
            
        except Exception as e:
            logger.error(f"Word 문서 로드 실패: {file_path}, 에러: {str(e)}")
            raise
    
    def load_excel(self, file_path: str, sheet_name: Optional[str] = None, **kwargs) -> List[Document]:
        """
        Excel 파일 로드
        
        Args:
            file_path: Excel 파일 경로
            sheet_name: 특정 시트 이름 (None이면 모든 시트)
            **kwargs: pandas read_excel에 전달할 추가 인자
        
        Returns:
            Document 객체 리스트 (시트별 또는 행별)
        """
        if not EXCEL_AVAILABLE:
            raise ImportError("Excel 처리를 위해 pandas와 openpyxl을 설치해주세요: pip install pandas openpyxl")
        
        documents = []
        file_name = Path(file_path).name
        
        try:
            # 모든 시트 이름 가져오기
            xl_file = pd.ExcelFile(file_path)
            
            # 처리할 시트 목록 결정
            if sheet_name:
                sheets_to_process = [sheet_name] if sheet_name in xl_file.sheet_names else []
            else:
                sheets_to_process = xl_file.sheet_names
            
            for sheet in sheets_to_process:
                df = pd.read_excel(file_path, sheet_name=sheet, **kwargs)
                
                # 시트 전체를 하나의 Document로
                sheet_text = f"Sheet: {sheet}\n\n"
                sheet_text += df.to_string(index=False)
                
                documents.append(Document(
                    page_content=sheet_text,
                    metadata={
                        'source': file_path,
                        'file_name': file_name,
                        'sheet_name': sheet,
                        'rows': len(df),
                        'columns': len(df.columns),
                        'column_names': list(df.columns),
                        'type': 'excel'
                    }
                ))
                
                # 옵션: 행별로도 추가 가능 (주석 해제하여 사용)
                # for idx, row in df.iterrows():
                #     row_text = ' | '.join([f"{col}: {val}" for col, val in row.items()])
                #     documents.append(Document(
                #         page_content=row_text,
                #         metadata={
                #             'source': file_path,
                #             'file_name': file_name,
                #             'sheet_name': sheet,
                #             'row': idx,
                #             'type': 'excel_row'
                #         }
                #     ))
            
            logger.info(f"Excel 파일 로드 완료: {file_name} ({len(documents)} 시트)")
            return documents
            
        except Exception as e:
            logger.error(f"Excel 파일 로드 실패: {file_path}, 에러: {str(e)}")
            raise
    
    def load_text(self, file_path: str, encoding: str = 'utf-8', **kwargs) -> List[Document]:
        """
        텍스트 파일 로드
        
        Args:
            file_path: 텍스트 파일 경로
            encoding: 파일 인코딩 (기본: utf-8)
            **kwargs: 추가 옵션
        
        Returns:
            Document 객체 리스트
        """
        file_name = Path(file_path).name
        
        try:
            with open(file_path, 'r', encoding=encoding) as file:
                content = file.read()
            
            documents = [Document(
                page_content=content,
                metadata={
                    'source': file_path,
                    'file_name': file_name,
                    'type': 'text',
                    'encoding': encoding
                }
            )]
            
            logger.info(f"텍스트 파일 로드 완료: {file_name}")
            return documents
            
        except Exception as e:
            logger.error(f"텍스트 파일 로드 실패: {file_path}, 에러: {str(e)}")
            raise
    
    def load_directory(self, directory_path: str, recursive: bool = False, 
                       file_extensions: Optional[List[str]] = None) -> List[Document]:
        """
        디렉토리 내의 모든 지원되는 파일 로드
        
        Args:
            directory_path: 디렉토리 경로
            recursive: 하위 디렉토리도 포함할지 여부
            file_extensions: 로드할 파일 확장자 리스트 (None이면 모두)
        
        Returns:
            Document 객체 리스트
        """
        directory = Path(directory_path)
        
        if not directory.exists() or not directory.is_dir():
            raise ValueError(f"유효한 디렉토리가 아닙니다: {directory_path}")
        
        all_documents = []
        
        # 파일 패턴 설정
        if recursive:
            pattern = '**/*'
        else:
            pattern = '*'
        
        # 지원 확장자 필터
        if file_extensions:
            extensions = [ext.lower() if ext.startswith('.') else f'.{ext.lower()}' 
                         for ext in file_extensions]
        else:
            extensions = list(self.supported_formats.keys())
        
        # 파일 로드
        for file_path in directory.glob(pattern):
            if file_path.is_file() and file_path.suffix.lower() in extensions:
                try:
                    documents = self.load(str(file_path))
                    all_documents.extend(documents)
                except Exception as e:
                    logger.warning(f"파일 로드 실패 (건너뜀): {file_path}, 에러: {str(e)}")
                    continue
        
        logger.info(f"디렉토리 로드 완료: {directory_path} ({len(all_documents)} 문서)")
        return all_documents


def split_documents(documents: List[Document], chunk_size: int = 500, 
                    chunk_overlap: int = 50) -> List[Document]:
    """
    문서를 작은 청크로 분할
    
    Args:
        documents: Document 객체 리스트
        chunk_size: 청크 크기 (문자 수)
        chunk_overlap: 청크 간 중첩 크기
    
    Returns:
        분할된 Document 객체 리스트
    """
    chunked_documents = []
    
    for doc in documents:
        text = doc.page_content
        
        # 텍스트가 chunk_size보다 작으면 그대로 유지
        if len(text) <= chunk_size:
            chunked_documents.append(doc)
            continue
        
        # 청크로 분할
        start = 0
        chunk_num = 0
        
        while start < len(text):
            end = start + chunk_size
            chunk_text = text[start:end]
            
            # 메타데이터 복사 및 청크 정보 추가
            chunk_metadata = doc.metadata.copy()
            chunk_metadata['chunk'] = chunk_num
            chunk_metadata['chunk_start'] = start
            chunk_metadata['chunk_end'] = end
            
            chunked_documents.append(Document(
                page_content=chunk_text,
                metadata=chunk_metadata
            ))
            
            # 다음 청크 시작 위치 (중첩 고려)
            start = end - chunk_overlap
            chunk_num += 1
    
    logger.info(f"문서 분할 완료: {len(documents)} → {len(chunked_documents)} 청크")
    return chunked_documents


# 사용 예제
if __name__ == '__main__':
    loader = DocumentLoader()
    
    # PDF 로드 예제
    # documents = loader.load_pdf('sample.pdf')
    
    # Word 문서 로드 예제
    # documents = loader.load_docx('sample.docx')
    
    # Excel 로드 예제
    # documents = loader.load_excel('sample.xlsx')
    
    # 디렉토리 로드 예제
    # documents = loader.load_directory('./documents', recursive=True)
    
    # 문서 분할 예제
    # chunked_docs = split_documents(documents, chunk_size=500, chunk_overlap=50)
    
    print("DocumentLoader 준비 완료")
